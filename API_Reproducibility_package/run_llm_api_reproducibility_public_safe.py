#!/usr/bin/env python3
"""
Public-safe reproducibility script for API-based execution of four LLM systems.

This script is designed for publication-facing reproducibility packages.
It does NOT store real API keys in the source code.
Instead, it reads secrets from environment variables or from a local .env file.

Supported systems:
- ChatGPT via OpenAI API
- Gemini via Google Gemini API
- Copilot-equivalent condition via Azure OpenAI API
- Perplexity via Perplexity API

Input structure:
01_question_sets/
    Group1_molecules/
    Group2_drugs/
    Group3_comparisons/
    Group4_inference/

Output structure:
02_api_answers/
    Group1_molecules/
        alanine_ChatGPT.docx
        alanine_Gemini.docx
        alanine_Copilot.docx
        alanine_Perplexity.docx
    ...

Important note for Group4:
The drug name must NOT be leaked to the model. The file name is used only for
local organization and output naming. The actual prompt sent to the model is
based only on the body text extracted from the Word document.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import time
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Callable, Dict, List, Optional, Sequence

import requests
from docx import Document

try:
    from dotenv import load_dotenv
except ImportError:  # pragma: no cover
    def load_dotenv(*args, **kwargs):
        return False

# Load local environment variables early.
SCRIPT_DIR = Path(__file__).resolve().parent
load_dotenv(dotenv_path=Path.cwd() / ".env", override=False)
if SCRIPT_DIR != Path.cwd():
    load_dotenv(dotenv_path=SCRIPT_DIR / ".env", override=False)

# =========================
# Constant configuration
# =========================
OPENAI_MODEL_ID = "gpt-4o"
GEMINI_MODEL_ID = "gemini-2.0-flash"
AZURE_TARGET_MODEL_FAMILY = "GPT-4 Turbo"
PERPLEXITY_MODEL_ID = "sonar"

DEFAULT_SYSTEM_INSTRUCTION = (
    "You are participating in an academic reproducibility workflow. "
    "Answer the user question directly and completely. "
    "Do not mention hidden file names, local paths, or implementation details."
)

DEFAULT_REQUEST_TIMEOUT_SECONDS = 180
DEFAULT_MAX_RETRIES = 5
DEFAULT_RETRY_BACKOFF_SECONDS = 2.0
DEFAULT_SLEEP_BETWEEN_CALLS_SECONDS = 1.0
RETRIABLE_STATUS_CODES = {408, 409, 425, 429, 500, 502, 503, 504}


# =========================
# Data structures
# =========================
@dataclass
class ModelConfig:
    public_name: str
    model_identifier: str
    caller: Callable[[str, str], str]


# =========================
# Environment helpers
# =========================
def get_env(name: str, default: str = "") -> str:
    return os.getenv(name, default).strip()


# =========================
# Utility functions
# =========================
def require_nonempty(value: str, label: str) -> None:
    if not value:
        raise ValueError(f"Missing required configuration: {label}")


def ensure_directory(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def sanitize_filename(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._()\- ]+", "_", name).strip()


def normalize_token(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", text.lower())


def clean_blank_lines(text: str) -> str:
    lines = [line.rstrip() for line in text.splitlines()]
    cleaned: List[str] = []
    previous_blank = False
    for line in lines:
        is_blank = not line.strip()
        if is_blank and previous_blank:
            continue
        cleaned.append(line)
        previous_blank = is_blank
    return "\n".join(cleaned).strip()


def read_docx_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    blocks: List[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                blocks.append("\t".join(cells))

    return "\n".join(blocks).strip()


def infer_group_name(docx_path: Path) -> str:
    parent = docx_path.parent.name.strip()
    return parent if parent else "Ungrouped"


def build_group4_name_variants(stem: str) -> List[str]:
    variants = set()
    base = stem.replace("_", " ").strip()
    if base:
        variants.add(base)

    without_parentheses = re.sub(r"\([^)]*\)", "", base).strip(" -_/,")
    if without_parentheses:
        variants.add(without_parentheses)

    for item in re.findall(r"\(([^)]*)\)", base):
        cleaned = item.strip(" -_/,")
        if cleaned:
            variants.add(cleaned)

    for item in re.split(r"[/,;]", base):
        cleaned = item.strip(" -_/,")
        if cleaned:
            variants.add(cleaned)

    return sorted(variants, key=len, reverse=True)


def anonymize_group4_text(text: str, stem: str) -> str:
    variants = build_group4_name_variants(stem)
    normalized_variants = {normalize_token(v) for v in variants if normalize_token(v)}

    filtered_lines: List[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped and normalize_token(stripped) in normalized_variants:
            continue
        filtered_lines.append(line)

    result = "\n".join(filtered_lines)

    for variant in variants:
        if len(normalize_token(variant)) < 3:
            continue
        result = re.sub(re.escape(variant), "", result, flags=re.IGNORECASE)

    return clean_blank_lines(result)


def prepare_prompt_for_group(docx_path: Path, raw_text: str) -> str:
    group_name = infer_group_name(docx_path).lower()

    if "group4" in group_name or "inference" in group_name:
        return anonymize_group4_text(raw_text, docx_path.stem)

    return clean_blank_lines(raw_text)


def parse_retry_after_seconds(value: str) -> Optional[float]:
    if not value:
        return None
    try:
        return max(0.0, float(value))
    except ValueError:
        return None


def post_json_with_retry(
    url: str,
    headers: Dict[str, str],
    payload: Dict,
    timeout: int = DEFAULT_REQUEST_TIMEOUT_SECONDS,
    max_retries: int = DEFAULT_MAX_RETRIES,
    retry_backoff_seconds: float = DEFAULT_RETRY_BACKOFF_SECONDS,
) -> Dict:
    last_error: Optional[BaseException] = None

    for attempt in range(1, max_retries + 1):
        try:
            response = requests.post(url, headers=headers, json=payload, timeout=timeout)

            if response.status_code in RETRIABLE_STATUS_CODES:
                response_text = response.text.strip()
                if attempt == max_retries:
                    raise RuntimeError(
                        f"HTTP {response.status_code} from {url}. Response body: {response_text[:1000]}"
                    )

                retry_after = parse_retry_after_seconds(response.headers.get("Retry-After", ""))
                sleep_seconds = retry_after if retry_after is not None else retry_backoff_seconds * (2 ** (attempt - 1))
                time.sleep(min(sleep_seconds, 60.0))
                continue

            response.raise_for_status()
            return response.json()
        except requests.RequestException as exc:
            last_error = exc
            if attempt == max_retries:
                raise RuntimeError(f"Request failed after {max_retries} attempts: {exc}") from exc
            time.sleep(min(retry_backoff_seconds * (2 ** (attempt - 1)), 60.0))
        except ValueError as exc:
            last_error = exc
            raise RuntimeError(f"The API response from {url} was not valid JSON.") from exc

    raise RuntimeError(f"Request failed after {max_retries} attempts: {last_error}")


def extract_openai_chat_text(data: Dict, provider_label: str) -> str:
    choices = data.get("choices", [])
    if not choices:
        raise RuntimeError(f"{provider_label} returned no choices.")

    message = choices[0].get("message", {})
    content = message.get("content", "")

    if isinstance(content, str):
        text = content.strip()
    elif isinstance(content, list):
        text_parts = [item.get("text", "") for item in content if isinstance(item, dict) and item.get("text")]
        text = "\n".join(text_parts).strip()
    else:
        text = ""

    if not text:
        raise RuntimeError(f"{provider_label} returned an empty text response.")

    return text


def extract_gemini_text(data: Dict) -> str:
    prompt_feedback = data.get("promptFeedback", {})
    block_reason = prompt_feedback.get("blockReason")
    if block_reason and not data.get("candidates"):
        raise RuntimeError(f"Gemini blocked the prompt. blockReason={block_reason}")

    candidates = data.get("candidates", [])
    if not candidates:
        raise RuntimeError("Gemini returned no candidates.")

    first_candidate = candidates[0]
    finish_reason = first_candidate.get("finishReason")
    parts = first_candidate.get("content", {}).get("parts", [])
    text_parts = [part.get("text", "") for part in parts if isinstance(part, dict) and part.get("text")]
    text = "\n".join(text_parts).strip()

    if text:
        return text

    if finish_reason == "SAFETY":
        raise RuntimeError("Gemini blocked the response for safety reasons.")

    raise RuntimeError(f"Gemini returned no text output. finishReason={finish_reason!r}")


def output_folder_for_file(docx_path: Path, input_root: Path, output_root: Path) -> Path:
    if docx_path.is_relative_to(input_root):
        relative_parent = docx_path.parent.relative_to(input_root)
        destination_folder = output_root / relative_parent
    else:
        destination_folder = output_root / "SingleFile"
    ensure_directory(destination_folder)
    return destination_folder


def write_answer_docx(
    output_path: Path,
    source_doc_name: str,
    source_group_name: str,
    model_public_name: str,
    model_identifier: str,
    prompt_text: str,
    answer_text: str,
    access_date: str,
    script_version: str,
    execution_timestamp: str,
) -> None:
    doc = Document()
    doc.add_heading(f"Response file for {source_doc_name}", level=1)
    doc.add_paragraph(f"Source group: {source_group_name}")
    doc.add_paragraph(f"Model label: {model_public_name}")
    doc.add_paragraph(f"Model identifier: {model_identifier}")
    doc.add_paragraph(f"Access date: {access_date}")
    doc.add_paragraph(f"Execution timestamp: {execution_timestamp}")
    doc.add_paragraph(f"Script version: {script_version}")

    doc.add_heading("Prompt", level=2)
    doc.add_paragraph(prompt_text)

    doc.add_heading("Raw output", level=2)
    doc.add_paragraph(answer_text)

    doc.save(str(output_path))


def write_json(path: Path, payload: Dict) -> None:
    with path.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, indent=2, ensure_ascii=False)


# =========================
# API caller implementations
# =========================
def call_openai(system_instruction: str, user_prompt: str) -> str:
    openai_api_key = get_env("OPENAI_API_KEY")
    require_nonempty(openai_api_key, "OPENAI_API_KEY")

    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {openai_api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": OPENAI_MODEL_ID,
        "messages": [
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": user_prompt},
        ],
    }
    data = post_json_with_retry(url=url, headers=headers, payload=payload)
    return extract_openai_chat_text(data, provider_label="OpenAI")


def call_gemini(system_instruction: str, user_prompt: str) -> str:
    gemini_api_key = get_env("GEMINI_API_KEY")
    require_nonempty(gemini_api_key, "GEMINI_API_KEY")

    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL_ID}:generateContent"
    headers = {
        "x-goog-api-key": gemini_api_key,
        "Content-Type": "application/json",
    }
    payload = {
        "system_instruction": {"parts": [{"text": system_instruction}]},
        "contents": [{"parts": [{"text": user_prompt}]}],
    }
    data = post_json_with_retry(url=url, headers=headers, payload=payload)
    return extract_gemini_text(data)


def call_azure(system_instruction: str, user_prompt: str) -> str:
    azure_api_key = get_env("AZURE_OPENAI_API_KEY")
    azure_endpoint = get_env("AZURE_OPENAI_ENDPOINT")
    azure_deployment = get_env("AZURE_OPENAI_DEPLOYMENT")

    require_nonempty(azure_api_key, "AZURE_OPENAI_API_KEY")
    require_nonempty(azure_endpoint, "AZURE_OPENAI_ENDPOINT")
    require_nonempty(azure_deployment, "AZURE_OPENAI_DEPLOYMENT")

    normalized_endpoint = azure_endpoint.rstrip("/")
    if not normalized_endpoint.endswith("/openai/v1"):
        normalized_endpoint = f"{normalized_endpoint}/openai/v1"

    url = f"{normalized_endpoint}/chat/completions"
    headers = {
        "api-key": azure_api_key,
        "Content-Type": "application/json",
    }
    payload = {
        "model": azure_deployment,
        "messages": [
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": user_prompt},
        ],
    }
    data = post_json_with_retry(url=url, headers=headers, payload=payload)
    return extract_openai_chat_text(data, provider_label="Azure OpenAI")


def call_perplexity(system_instruction: str, user_prompt: str) -> str:
    perplexity_api_key = get_env("PERPLEXITY_API_KEY")
    require_nonempty(perplexity_api_key, "PERPLEXITY_API_KEY")

    url = "https://api.perplexity.ai/chat/completions"
    headers = {
        "Authorization": f"Bearer {perplexity_api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": PERPLEXITY_MODEL_ID,
        "messages": [
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": user_prompt},
        ],
    }
    data = post_json_with_retry(url=url, headers=headers, payload=payload)
    return extract_openai_chat_text(data, provider_label="Perplexity")


# =========================
# Main processing logic
# =========================
def azure_model_identifier() -> str:
    deployment = get_env("AZURE_OPENAI_DEPLOYMENT")
    if deployment:
        return f"Azure deployment: {deployment}; target family: {AZURE_TARGET_MODEL_FAMILY}"
    return f"Azure deployment: <unset>; target family: {AZURE_TARGET_MODEL_FAMILY}"


def build_model_registry(selected_models: Optional[Sequence[str]]) -> List[ModelConfig]:
    selected = {item.lower() for item in selected_models} if selected_models else None

    registry = [
        ModelConfig("ChatGPT", OPENAI_MODEL_ID, call_openai),
        ModelConfig("Gemini", GEMINI_MODEL_ID, call_gemini),
        ModelConfig("Copilot", azure_model_identifier(), call_azure),
        ModelConfig("Perplexity", PERPLEXITY_MODEL_ID, call_perplexity),
    ]

    if selected is None:
        return registry

    filtered = [model for model in registry if model.public_name.lower() in selected]
    if not filtered:
        raise ValueError("No valid models were selected. Use names such as ChatGPT, Gemini, Copilot, or Perplexity.")
    return filtered


def collect_docx_files(input_root: Path, single_file: Optional[Path]) -> List[Path]:
    if single_file is not None:
        if single_file.suffix.lower() != ".docx":
            raise ValueError(f"The --single-file argument must point to a .docx file: {single_file}")
        return [single_file]

    return sorted(path for path in input_root.rglob("*.docx") if not path.name.startswith("~$"))


def process_one_file(
    docx_path: Path,
    input_root: Path,
    output_root: Path,
    model_registry: Sequence[ModelConfig],
    system_instruction: str,
    access_date: str,
    script_version: str,
    dry_run: bool,
    sleep_between_calls_seconds: float,
) -> None:
    raw_text = read_docx_text(docx_path)
    if not raw_text:
        raise RuntimeError(f"The Word file is empty or unreadable: {docx_path}")

    prompt_text = prepare_prompt_for_group(docx_path, raw_text)
    if not prompt_text:
        raise RuntimeError(f"The prepared prompt became empty after preprocessing: {docx_path}")

    destination_folder = output_folder_for_file(docx_path, input_root, output_root)
    stem = sanitize_filename(docx_path.stem)
    group_name = infer_group_name(docx_path)

    for model in model_registry:
        output_name = f"{stem}_{model.public_name}.docx"
        output_path = destination_folder / output_name
        execution_timestamp = datetime.now().isoformat(timespec="seconds")

        if dry_run:
            answer_text = f"DRY RUN: no API call was made for {model.public_name}."
        else:
            answer_text = model.caller(system_instruction, prompt_text)
            if sleep_between_calls_seconds > 0:
                time.sleep(sleep_between_calls_seconds)

        write_answer_docx(
            output_path=output_path,
            source_doc_name=docx_path.name,
            source_group_name=group_name,
            model_public_name=model.public_name,
            model_identifier=model.model_identifier,
            prompt_text=prompt_text,
            answer_text=answer_text,
            access_date=access_date,
            script_version=script_version,
            execution_timestamp=execution_timestamp,
        )


def build_metadata_payload(
    input_root: Path,
    output_root: Path,
    selected_models: Sequence[ModelConfig],
    access_date: str,
    script_version: str,
    processed_files: Sequence[Path],
    single_file: Optional[Path],
    dry_run: bool,
) -> Dict:
    return {
        "package_type": "API-based reproducibility supplement",
        "script_version": script_version,
        "access_date": access_date,
        "run_timestamp": datetime.now().isoformat(timespec="seconds"),
        "input_root": str(input_root),
        "output_root": str(output_root),
        "single_file": str(single_file) if single_file else None,
        "dry_run": dry_run,
        "model_settings": [
            {
                "public_name": model.public_name,
                "model_identifier": model.model_identifier,
            }
            for model in selected_models
        ],
        "processed_files": [str(path) for path in processed_files],
        "notes": [
            "This script is intended for reproducibility supplementation.",
            "Real API keys must not be embedded in the public repository.",
            "Group4 prompts are anonymized before transmission to the model.",
        ],
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run the API-based LLM reproducibility workflow.")
    parser.add_argument(
        "--input-root",
        default=get_env("INPUT_ROOT"),
        help="Path to the root question-set folder. If omitted, INPUT_ROOT from .env is used.",
    )
    parser.add_argument(
        "--output-root",
        default=get_env("OUTPUT_ROOT"),
        help="Path to the root output folder. If omitted, OUTPUT_ROOT from .env is used.",
    )
    parser.add_argument(
        "--single-file",
        default="",
        help="Optional: run only one specific .docx file instead of all files under input-root.",
    )
    parser.add_argument(
        "--models",
        nargs="*",
        default=None,
        help="Optional subset of models to run. Examples: ChatGPT Gemini Copilot Perplexity",
    )
    parser.add_argument(
        "--access-date",
        default=get_env("ACCESS_DATE", date.today().isoformat()),
        help="Date string to store in the generated files. If omitted, today's date is used.",
    )
    parser.add_argument(
        "--system-instruction",
        default=DEFAULT_SYSTEM_INSTRUCTION,
        help="Optional system instruction to prepend to every API call.",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Create output files without calling any API. Useful for validating folder logic.",
    )
    parser.add_argument(
        "--script-version",
        default="1.1.0",
        help="Version string to store in metadata output.",
    )
    parser.add_argument(
        "--sleep-between-calls",
        type=float,
        default=float(get_env("SLEEP_BETWEEN_CALLS", str(DEFAULT_SLEEP_BETWEEN_CALLS_SECONDS))),
        help="Delay in seconds between API calls. If omitted, SLEEP_BETWEEN_CALLS or 1.0 is used.",
    )
    return parser.parse_args()


def validate_paths(input_root: Path, output_root: Path, single_file: Optional[Path]) -> None:
    if not input_root.exists():
        raise FileNotFoundError(f"Input root does not exist: {input_root}")
    if not input_root.is_dir():
        raise NotADirectoryError(f"Input root is not a directory: {input_root}")
    if single_file is not None and not single_file.exists():
        raise FileNotFoundError(f"Single file does not exist: {single_file}")
    if single_file is not None and single_file.suffix.lower() != ".docx":
        raise ValueError(f"Single file must be a .docx file: {single_file}")
    ensure_directory(output_root)


def main() -> int:
    args = parse_args()

    if not args.input_root:
        raise ValueError("Missing input root. Set INPUT_ROOT in .env or pass --input-root.")
    if not args.output_root:
        raise ValueError("Missing output root. Set OUTPUT_ROOT in .env or pass --output-root.")

    input_root = Path(args.input_root).expanduser().resolve()
    output_root = Path(args.output_root).expanduser().resolve()
    single_file = Path(args.single_file).expanduser().resolve() if args.single_file else None

    validate_paths(input_root=input_root, output_root=output_root, single_file=single_file)

    model_registry = build_model_registry(args.models)
    docx_files = collect_docx_files(input_root=input_root, single_file=single_file)
    if not docx_files:
        raise FileNotFoundError("No .docx files were found for processing.")

    for docx_path in docx_files:
        process_one_file(
            docx_path=docx_path,
            input_root=input_root,
            output_root=output_root,
            model_registry=model_registry,
            system_instruction=args.system_instruction,
            access_date=args.access_date,
            script_version=args.script_version,
            dry_run=args.dry_run,
            sleep_between_calls_seconds=args.sleep_between_calls,
        )

    metadata_path = output_root / "run_metadata.json"
    metadata_payload = build_metadata_payload(
        input_root=input_root,
        output_root=output_root,
        selected_models=model_registry,
        access_date=args.access_date,
        script_version=args.script_version,
        processed_files=docx_files,
        single_file=single_file,
        dry_run=args.dry_run,
    )
    write_json(metadata_path, metadata_payload)

    print(f"Processed {len(docx_files)} file(s).")
    print(f"Outputs written to: {output_root}")
    print(f"Metadata written to: {metadata_path}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise
